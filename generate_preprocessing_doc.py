from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.name = "Calibri"
    if level == 1:
        h.font.size = Pt(18)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(14)
        h.font.bold = True
    else:
        h.font.size = Pt(12)
        h.font.bold = True

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h_text in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        shading_elm = cell._element.get_or_add_tcPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "D9D9D9")
        shade.set(qn("w:val"), "clear")
        shading_elm.append(shade)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(11)

# ── TITLE ──
doc.add_heading("Data Preprocessing Documentation", level=1)

p = doc.add_paragraph()
run = p.add_run("Dataset: Coderush-26 ML Module")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
run = p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
run = p.add_run("Script: preprocess.py")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
run = p.add_run("Input: Coderush-26-ML-Train.csv, Coderush-26-ML-test.csv")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
run = p.add_run("Output: processed_data/ (X_train.csv, y_train.csv, X_test.csv, train_bag_ids.npy, test_bag_ids.npy, preprocessing_artifacts.pkl)")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()

# ═══════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════
doc.add_heading("1. Preprocessing Pipeline Overview", level=1)

doc.add_paragraph(
    "This document describes the data preprocessing pipeline applied to the Coderush-26 ML dataset "
    "in preparation for model training. All preprocessing decisions are based on findings from the "
    "Exploratory Data Analysis (EDA) phase. The pipeline transforms raw data with 29 columns into "
    "a clean feature matrix of 54 columns, ready for model training."
)

add_table(doc,
    ["Stage", "Description", "Output"],
    [
        ["1. Drop Columns", "Remove leakage, noise, redundant features", "22 columns"],
        ["2. Bag Features", "Aggregate household-level statistics", "+21 features"],
        ["3. Individual Features", "Derive age, capital ratios, log transforms", "+12 features"],
        ["4. Encode Categoricals", "Ordinal, label, and target encoding", "All numeric"],
        ["5. Clip Outliers", "Cap extreme values at 99.5th percentile", "Bounded features"],
        ["6. Finalize", "Align train/test columns, fill NaNs", "54 features"],
    ]
)
doc.add_paragraph()

# ═══════════════════════════════════════════
# 2. FEATURE DROPS
# ═══════════════════════════════════════════
doc.add_heading("2. Dropped Features", level=1)

doc.add_paragraph(
    "Eight features were removed from the dataset. Each drop is justified by EDA findings "
    "that identified the feature as constant, near-constant, a data identifier, or redundant."
)

add_table(doc,
    ["Feature", "Reason", "EDA Evidence"],
    [
        ["processing_flag", "Constant", "All 16,776 rows = 1.0 (100%)"],
        ["survey_year", "Constant", "All 16,776 rows = 1994 (100%)"],
        ["currency_code", "Constant", "All 16,776 rows = 'USD' (100%)"],
        ["poverty_line_usd", "Constant", "All 16,776 rows = 15141 (100%)"],
        ["is_adult_flag", "Near-constant", "16,669 rows = 1, only 107 rows = 0 (99.4% one class)"],
        ["person_idx", "Identifier", "Position index within bag (0-7), no predictive signal"],
        ["annual_hours_est", "Redundant", "Perfectly correlated with hours_per_week (derived from it)"],
        ["bag_id", "Identifier", "Household ID — dropped after bag-level aggregation"],
    ]
)
doc.add_paragraph()

doc.add_paragraph("Impact:", style="Heading 3")
add_bullet(doc, "Reduced feature count from 29 to 22 before engineering.")
add_bullet(doc, "Eliminated risk of models learning spurious patterns from constants.")
add_bullet(doc, "Removed redundant feature that would add no information but increase dimensionality.")

# ═══════════════════════════════════════════
# 3. BAG-LEVEL FEATURE ENGINEERING
# ═══════════════════════════════════════════
doc.add_heading("3. Bag-Level Feature Engineering", level=1)

doc.add_paragraph(
    "Each person belongs to a household (bag) of 3-7 members. EDA revealed that all members "
    "within a bag share the same label, meaning household-level characteristics are predictive. "
    "We compute 21 aggregate statistics per bag and merge them back to each individual row."
)

doc.add_heading("3.1 Bag Composition Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_member_count", "First value of bag_size", "Household size affects economic dynamics"],
        ["bag_size", "First value of bag_size", "Redundant with member_count; kept for clarity"],
    ]
)
doc.add_paragraph()

doc.add_heading("3.2 Education Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_education_mean", "Mean of education_num", "Average education level in household"],
        ["bag_education_std", "Std of education_num", "Education diversity within household"],
        ["bag_education_range", "Max - Min of education_num", "Spread of education levels"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "Education was the strongest predictor in EDA (r=+0.144). Household-level education "
           "captures the collective human capital of the family.")
add_bullet(doc, "Range captures inequality — a household with mixed education levels may have "
           "different economic dynamics than one with uniform education.")

doc.add_heading("3.3 Age Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_age_mean", "Mean of year_of_birth", "Average age/generation of household"],
        ["bag_age_range", "Max - Min of year_of_birth", "Generational span (e.g., parents + adult children)"],
    ]
)
doc.add_paragraph()

doc.add_heading("3.4 Work Hours Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_hours_mean", "Mean of hours_per_week", "Average work intensity in household"],
        ["bag_hours_std", "Std of hours_per_week", "Variation in work patterns"],
        ["bag_hours_range", "Max - Min of hours_per_week", "Difference between hardest and least working member"],
    ]
)
doc.add_paragraph()

doc.add_heading("3.5 Capital Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_capital_gain_max", "Max of capital_gain", "Highest earner in household drives class"],
        ["bag_capital_gain_mean", "Mean of capital_gain", "Average capital income across household"],
        ["bag_capital_loss_max", "Max of capital_loss", "Maximum loss experienced"],
        ["bag_capital_loss_mean", "Mean of capital_loss", "Average loss across household"],
        ["bag_net_capital_max", "Max of net_capital_asset", "Best net position in household"],
        ["bag_net_capital_mean", "Mean of net_capital_asset", "Average net capital across household"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "Capital features showed the clearest class gradient in EDA. Household-level "
           "capital aggregates amplify this signal.")
add_bullet(doc, "Using both max and mean captures two different dynamics: the 'top earner' effect "
           "and the 'collective wealth' effect.")

doc.add_heading("3.6 Diversity and Activity Features", level=2)

add_table(doc,
    ["Feature", "Aggregation", "Rationale"],
    [
        ["bag_duration_mean", "Mean of survey_duration_mins", "Average survey completion time"],
        ["bag_unique_relationships", "Count of unique relationships", "Household structure complexity"],
        ["bag_unique_occupations", "Count of unique occupations", "Occupational diversity"],
        ["bag_unique_workclass", "Count of unique workclass", "Employment type diversity"],
        ["bag_has_capital_activity", "Max of capital_activity_flag", "Whether any member has capital activity"],
    ]
)
doc.add_paragraph()

# ═══════════════════════════════════════════
# 4. INDIVIDUAL-LEVEL FEATURE ENGINEERING
# ═══════════════════════════════════════════
doc.add_heading("4. Individual-Level Feature Engineering", level=1)

doc.add_paragraph(
    "Beyond bag-level aggregates, we create 12 individual-level derived features that capture "
    "non-linear relationships and domain-specific economic signals."
)

doc.add_heading("4.1 Age Features", level=2)

add_table(doc,
    ["Feature", "Formula", "Rationale"],
    [
        ["age", "1994 - year_of_birth", "Convert birth year to actual age (survey year = 1994)"],
        ["age_squared", "age^2", "Capture non-linear age effect — income peaks mid-career"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "year_of_birth had r=-0.114 with target. Converting to age makes the relationship "
           "more interpretable.")
add_bullet(doc, "age_squared captures the inverted-U relationship between age and economic class "
           "(young and old tend toward lower classes, middle-aged toward upper).")

doc.add_heading("4.2 Capital Features", level=2)

add_table(doc,
    ["Feature", "Formula", "Rationale"],
    [
        ["net_capital_clean", "capital_gain - capital_loss", "Net capital position"],
        ["capital_ratio", "capital_gain / (capital_loss + 1)", "Gain-to-loss ratio; +1 prevents division by zero"],
        ["has_capital_gain", "capital_gain > 0 (binary)", "Whether person had any capital gains"],
        ["has_capital_loss", "capital_loss > 0 (binary)", "Whether person had any capital losses"],
        ["capital_per_hour", "capital_gain / (hours_per_week + 1)", "Capital efficiency per working hour"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "Capital features were 80-93% zero-inflated. Binary flags (has_capital_gain, "
           "has_capital_loss) capture the presence/absence signal separately from magnitude.")
add_bullet(doc, "capital_ratio distinguishes between people who gain more than they lose versus "
           "those who lose more than they gain.")
add_bullet(doc, "capital_per_hour normalizes capital by work effort — high capital with low hours "
           "suggests investment income rather than wages.")

doc.add_heading("4.3 Work Pattern Features", level=2)

add_table(doc,
    ["Feature", "Formula", "Rationale"],
    [
        ["works_full_time", "hours_per_week >= 35 (binary)", "Standard full-time threshold"],
        ["works_overtime", "hours_per_week > 40 (binary)", "Working beyond standard 40-hour week"],
    ]
)
doc.add_paragraph()

doc.add_heading("4.4 Interaction Features", level=2)

add_table(doc,
    ["Feature", "Formula", "Rationale"],
    [
        ["age_x_education", "age * education_num", "Interaction: education value may vary by life stage"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "A 25-year-old with a bachelor's degree has different prospects than a 55-year-old "
           "with the same degree. This interaction captures that effect.")

doc.add_heading("4.5 Log Transformations", level=2)

add_table(doc,
    ["Feature", "Formula", "Rationale"],
    [
        ["log_capital_gain", "log(1 + capital_gain)", "Compress right-skewed distribution"],
        ["log_capital_loss", "log(1 + capital_loss)", "Compress right-skewed distribution"],
        ["log_net_capital_asset", "log(1 + net_capital_asset)", "Compress right-skewed distribution"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "capital_gain: mean=2,049, max=99,999, 87% zeros. Log transformation reduces "
           "the influence of extreme outliers while preserving the zero-inflation structure.")
add_bullet(doc, "log1p is used (log of 1 + value) so that zero values map to 0 rather than negative infinity.")

# ═══════════════════════════════════════════
# 5. CATEGORICAL ENCODING
# ═══════════════════════════════════════════
doc.add_heading("5. Categorical Encoding Strategy", level=1)

doc.add_paragraph(
    "Three encoding methods are used based on feature cardinality and semantic ordering. "
    "The choice of encoding method is critical for model performance."
)

doc.add_heading("5.1 Ordinal Encoding", level=2)

doc.add_paragraph("Used for features with a natural order:")

add_table(doc,
    ["Feature", "Mapping", "Rationale"],
    [
        ["education_tier", "Primary=0, Secondary=1, Higher=2", "Clear educational progression"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "education_tier showed a clear gradient in EDA: Primary -> 21.7% upper, "
           "Secondary -> 29.8% upper, Higher -> 38.3% upper.")
add_bullet(doc, "Ordinal encoding preserves this ordering so the model can learn the monotonic relationship.")

doc.add_heading("5.2 Label Encoding", level=2)

doc.add_paragraph("Used for low-cardinality features without strong ordinal meaning:")

add_table(doc,
    ["Feature", "Unique Values", "Method"],
    [
        ["relationship", "~7", "Category codes (alphabetical order)"],
        ["marital_status", "~6", "Category codes (alphabetical order)"],
        ["sex", "2", "Category codes (Female=0, Male=1)"],
        ["interview_mode", "2", "Category codes"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "These features have few unique values, so arbitrary integer assignment is acceptable.")
add_bullet(doc, "Tree-based models (LightGBM, XGBoost) are invariant to the specific integer mapping "
           "for low-cardinality features.")

doc.add_heading("5.3 Target Encoding (with Smoothing)", level=2)

doc.add_paragraph("Used for high-cardinality features where label encoding would create sparse high-dimensional space:")

add_table(doc,
    ["Feature", "Unique Values", "Encoding"],
    [
        ["occupation", "~14", "Smoothed mean of encoded label"],
        ["workclass", "~8", "Smoothed mean of encoded label"],
        ["education", "~16", "Smoothed mean of encoded label"],
        ["native_country", "~41", "Smoothed mean of encoded label"],
        ["race", "~5", "Smoothed mean of encoded label"],
    ]
)
doc.add_paragraph()

doc.add_paragraph("Target encoding formula:", style="Heading 3")

p = doc.add_paragraph()
run = p.add_run("encoded_value = (category_mean * count + global_mean * smoothing) / (count + smoothing)")
run.font.name = "Courier New"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()

add_bullet(doc, "Label is encoded as: lower=0, middle=1, upper=2.")
add_bullet(doc, "Smoothing parameter k=10 prevents overfitting for rare categories.")
add_bullet(doc, "native_country has ~41 unique values, making it a prime candidate for target encoding "
           "rather than one-hot (which would add 41 columns).")
add_bullet(doc, "Original categorical columns are dropped after encoding to prevent data leakage "
           "and reduce dimensionality.")
add_bullet(doc, "Unseen categories in test data are filled with the global mean.")

# ═══════════════════════════════════════════
# 6. OUTLIER HANDLING
# ═══════════════════════════════════════════
doc.add_heading("6. Outlier Handling", level=1)

doc.add_paragraph(
    "Capital features contain extreme values that can disproportionately influence model training. "
    "We clip these at the 99.5th percentile, computed on the training data."
)

add_table(doc,
    ["Feature", "99.5th Percentile Clip Value", "Original Max"],
    [
        ["capital_gain", "99,999", "99,999"],
        ["capital_loss", "2,415", "3,900"],
        ["net_capital_asset", "99,999", "99,999"],
        ["net_capital_clean", "99,999", "99,999"],
        ["capital_ratio", "99,999", "~99,999"],
        ["capital_per_hour", "1,961", "~99,999"],
        ["bag_capital_gain_max", "99,999", "99,999"],
        ["bag_capital_gain_mean", "26,233", "~99,999"],
        ["bag_capital_loss_max", "2,824", "3,900"],
        ["bag_capital_loss_mean", "941", "~3,900"],
        ["bag_net_capital_max", "99,999", "99,999"],
        ["bag_net_capital_mean", "26,233", "~99,999"],
    ]
)
doc.add_paragraph()

add_bullet(doc, "Clipping is applied to both train and test sets using training percentiles only "
           "(prevents data leakage from test set statistics).")
add_bullet(doc, "Some values (capital_gain, net_capital_asset) already max at 99,999, suggesting "
           "possible censoring in the original data collection.")
add_bullet(doc, "Log-transformed versions (log_capital_gain, etc.) are not clipped as the log "
           "function already compresses the scale.")

# ═══════════════════════════════════════════
# 7. OUTPUT ARTIFACTS
# ═══════════════════════════════════════════
doc.add_heading("7. Output Artifacts", level=1)

doc.add_paragraph("The pipeline produces the following files in the processed_data/ directory:")

add_table(doc,
    ["File", "Format", "Contents", "Shape"],
    [
        ["X_train.csv", "CSV", "Feature matrix for training", "16,776 x 54"],
        ["y_train.csv", "CSV", "Target labels (encoded: 0=lower, 1=middle, 2=upper)", "16,776 x 1"],
        ["X_test.csv", "CSV", "Feature matrix for test predictions", "1,981 x 54"],
        ["train_bag_ids.npy", "NumPy", "Bag IDs for each training row (GroupKFold)", "16,776"],
        ["test_bag_ids.npy", "NumPy", "Bag IDs for each test row", "1,981"],
        ["preprocessing_artifacts.pkl", "Pickle", "Encoding maps, clip values, config", "Dictionary"],
    ]
)
doc.add_paragraph()

doc.add_heading("7.1 Preprocessing Artifacts Contents", level=2)

add_table(doc,
    ["Key", "Type", "Description"],
    [
        ["label_to_int", "dict", "Mapping: lower=0, middle=1, upper=2"],
        ["int_to_label", "dict", "Reverse mapping for prediction decoding"],
        ["ordinal_mappings", "dict", "Ordinal encoding maps (education_tier)"],
        ["target_encodings", "dict", "Target encoding value maps per category per feature"],
        ["drop_cols", "list", "List of columns dropped during preprocessing"],
        ["feature_cols", "list", "Final list of 54 feature column names"],
        ["clip_values", "dict", "99.5th percentile clip value per feature"],
    ]
)
doc.add_paragraph()

# ═══════════════════════════════════════════
# 8. FEATURE SUMMARY
# ═══════════════════════════════════════════
doc.add_heading("8. Final Feature Inventory", level=1)

doc.add_paragraph(
    "The final feature matrix contains 54 columns, organized into the following categories:"
)

add_table(doc,
    ["Category", "Count", "Examples"],
    [
        ["Original numerical", "~7", "education_num, hours_per_week, survey_duration_mins, capital_gain, capital_loss, net_capital_asset, bag_size"],
        ["Original encoded categorical", "~9", "relationship, marital_status, sex, interview_mode, education_tier, + 5 target-encoded features"],
        ["Bag-level aggregates", "21", "bag_education_mean, bag_capital_gain_max, bag_hours_std, etc."],
        ["Individual derived", "12", "age, age_squared, net_capital_clean, capital_ratio, log_capital_gain, etc."],
        ["Binary indicators", "~5", "has_capital_gain, has_capital_loss, works_full_time, works_overtime"],
    ]
)
doc.add_paragraph()

# ═══════════════════════════════════════════
# 9. TRAINING RECOMMENDATIONS
# ═══════════════════════════════════════════
doc.add_heading("9. Training Recommendations", level=1)

doc.add_heading("9.1 Cross-Validation Strategy", level=2)

add_bullet(doc, "Use GroupKFold with bag_id as the group key. Do NOT use random KFold or StratifiedKFold "
           "without groups, as this would leak household information between folds.")
add_bullet(doc, "Recommended: 5-fold GroupKFold. With 3,360 bags, each fold will have ~672 bags.")
add_bullet(doc, "Alternative: GroupStratifiedKFold if available, to preserve class distribution across folds.")

doc.add_heading("9.2 Model Selection", level=2)

add_bullet(doc, "Baseline: Logistic Regression with class_weight='balanced'. Establishes a performance floor.")
add_bullet(doc, "Primary: LightGBM with objective='multiclass', num_class=3, class_weight='balanced'. "
           "Handles mixed feature types well and is fast.")
add_bullet(doc, "Secondary: XGBoost with similar settings. Provides complementary predictions for ensembling.")
add_bullet(doc, "Ensemble: Average predictions from LightGBM + XGBoost for marginal gains.")

doc.add_heading("9.3 Hyperparameter Tuning Priorities", level=2)

add_bullet(doc, "LightGBM: num_leaves, learning_rate, n_estimators, min_child_samples, reg_alpha, reg_lambda")
add_bullet(doc, "Class weights: May need manual tuning beyond 'balanced' if minority class is still underperforming.")
add_bullet(doc, "Feature selection: Consider removing low-importance features after initial model to reduce overfitting.")

doc.add_heading("9.4 Evaluation Metric", level=2)

add_bullet(doc, "Primary metric: Macro F1 Score. Treats all classes equally regardless of size.")
add_bullet(doc, "Monitor per-class F1: Ensure 'lower' class (minority) is not being sacrificed.")
add_bullet(doc, "Also track: Confusion matrix to understand common misclassifications (e.g., lower <-> middle).")

# ═══════════════════════════════════════════
# 10. KNOWN LIMITATIONS
# ═══════════════════════════════════════════
doc.add_heading("10. Known Limitations and Risks", level=1)

add_bullet(doc, "Target encoding uses the full training set. For production, consider K-fold target encoding "
           "to prevent target leakage.")
add_bullet(doc, "Bag-level features assume test bags are similarly structured. If test bags differ "
           "significantly in composition, bag features may not generalize.")
add_bullet(doc, "Log transformation + clipping is applied independently. In some cases, clipping after "
           "log transform might be more appropriate.")
add_bullet(doc, "The smoothing parameter (k=10) for target encoding was chosen heuristically. "
           "Cross-validation may reveal a better value.")
add_bullet(doc, "No feature selection has been applied yet. Some of the 54 features may be redundant "
           "or noisy. Feature importance analysis after initial training is recommended.")

doc.save("documentation/Data_Preprocessing_Documentation.docx")
print("Saved: documentation/Data_Preprocessing_Documentation.docx")
